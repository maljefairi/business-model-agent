import requests
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def get_ollama_response(prompt):
    url = "http://localhost:11434/api/generate"
    data = {
        "model": "gemma2",
        "prompt": prompt,
        "stream": False
    }
    response = requests.post(url, json=data)
    
    try:
        json_response = response.json()
        return json_response.get('response', ''), json_response
    except requests.exceptions.JSONDecodeError as e:
        print("Error decoding JSON response:")
        print(f"Response status code: {response.status_code}")
        print(f"Response content: {response.text}")
        raise e

def generate_business_model(project_title, project_description):
    sections = [
        "Executive Summary",
        "Company Description",
        "Market Analysis",
        "Products and Services",
        "Marketing and Sales Strategy",
        "Financial Projections",
        "Funding Requirements",
        "Implementation Plan",
        "Risk Analysis",
        "Conclusion"
    ]

    business_model = {}
    for section in sections:
        prompt = f"As a global expert in business strategy and entrepreneurship, based on the project titled '{project_title}' with the description '{project_description}', generate the '{section}' section of a comprehensive business model. Provide detailed, specific, and insightful information for this section only, without repeating content from other sections. Your response should demonstrate deep industry knowledge, strategic thinking, and innovative approaches. Do not use any markdown formatting like '#' or '*' in your response."
        content, _ = get_ollama_response(prompt)
        business_model[section] = content

    return business_model

def save_to_docx(business_model, filename, project_title):
    doc = Document()
    
    # Add title page
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(project_title)
    title_run.font.size = Pt(48)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.space_after = Pt(0)
    
    # Add page break after title
    doc.add_page_break()
    
    # Add content
    for section, content in business_model.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    
    # Save the document
    doc.save(filename)

def main():
    project_title = "Asas LLC"
    project_description = "Islamic Shariah compliant auditing and accounting firm, providing a comprehensive range of services to meet the diverse needs of our clients including certifications on halal products and services, and training, and other services, based on global benchmarks."

    business_model = generate_business_model(project_title, project_description)
    
    filename = f"{project_title}1.docx"
    
    save_to_docx(business_model, filename, project_title)
    print(f"Business model saved to {filename}")

if __name__ == "__main__":
    main()